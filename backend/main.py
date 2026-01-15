from fileinput import filename
import os
import time
import io
import json
from typing import List, Optional
from fastapi import FastAPI, HTTPException, UploadFile, File
from pydantic import BaseModel
from dotenv import load_dotenv
from pinecone import Pinecone, ServerlessSpec
from google import genai
from fastapi.middleware.cors import CORSMiddleware
from pypdf import PdfReader
from docx import Document

# ==========================================
# 1. SETUP & CONFIGURATION
# ==========================================
load_dotenv()

# Load API Keys with safety checks
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME")
KNOWLEDGE_MAP_INDEX_NAME = os.getenv("KNOWLEDGE_MAP_INDEX_NAME")

if not all([GEMINI_API_KEY, PINECONE_API_KEY, PINECONE_INDEX_NAME, KNOWLEDGE_MAP_INDEX_NAME]):
    raise ValueError("❌ Missing required API keys in .env")

# Configure Gemini with new API
client = genai.Client(api_key=GEMINI_API_KEY)

# --- CHANGED: Use 2.5-flash-lite ---
CHAT_MODEL_NAME = 'gemini-2.5-flash-lite'
EMBED_MODEL_NAME = 'text-embedding-004'

# Pinecone Configuration
pc = Pinecone(api_key=PINECONE_API_KEY)

# Main Knowledge Base Index
if PINECONE_INDEX_NAME not in pc.list_indexes().names():
    print(f"⚙️ Creating Knowledge Base index '{PINECONE_INDEX_NAME}'")
    pc.create_index(
        name=PINECONE_INDEX_NAME,
        dimension=768,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1")
    )
index_kb = pc.Index(PINECONE_INDEX_NAME)

# Knowledge Map Index
if KNOWLEDGE_MAP_INDEX_NAME not in pc.list_indexes().names():
    print(f"⚙️ Creating Knowledge Map index '{KNOWLEDGE_MAP_INDEX_NAME}'")
    pc.create_index(
        name=KNOWLEDGE_MAP_INDEX_NAME,
        dimension=768,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region="us-east-1")
    )
index_km = pc.Index(KNOWLEDGE_MAP_INDEX_NAME)

# ==========================================
# OPTIMIZED CHUNKING PARAMETERS
# ==========================================
# Recommended for text-embedding-004 (768 dims) + conversational AI
CHUNK_SIZE = 512              # Tokens ≈ 400-450 words - optimal for semantic coherence
CHUNK_OVERLAP = 128           # 25% overlap prevents context loss at boundaries
BATCH_SIZE = 100              # Pinecone upsert batch size

# Character-based approximation (1 token ≈ 4 characters for English)
CHUNK_SIZE_CHARS = CHUNK_SIZE * 4      # ~2048 characters
CHUNK_OVERLAP_CHARS = CHUNK_OVERLAP * 4 # ~512 characters

# ==========================================
# 2. DATA MODELS (Pydantic)
# ==========================================
class ChatRequest(BaseModel):
    query: str    
    top_k: Optional[int] = 5  # Increased from 3 for better context coverage
    system_prompt: Optional[str] = None

class Source(BaseModel):
    text: str
    source: str
    score: float

class ChatResponse(BaseModel):
    response: str          
    sources: List[Source]

class IngestResponse(BaseModel):
    success: bool
    message: str
    chunks_added: int
    filename: str

# ==========================================
# 3. IMPROVED CHUNKING FUNCTIONS
# ==========================================
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes"""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        return "".join([page.extract_text() or "" for page in reader.pages])
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes"""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def smart_chunk_text(text: str, chunk_size: int = CHUNK_SIZE_CHARS, 
                     overlap: int = CHUNK_OVERLAP_CHARS) -> List[str]:
    """
    Intelligent text chunking that respects sentence boundaries.
    Better than naive splitting for maintaining semantic coherence.
    """
    # Split on sentence boundaries
    sentences = text.replace('\n', ' ').split('. ')
    sentences = [s.strip() + '.' for s in sentences if s.strip()]
    
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        sentence_len = len(sentence)
        
        # If single sentence exceeds chunk size, split it
        if sentence_len > chunk_size:
            if current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_size = 0
            
            # Split long sentence by character limit
            for i in range(0, len(sentence), chunk_size - overlap):
                chunks.append(sentence[i:i + chunk_size])
            continue
        
        # Add sentence to current chunk
        if current_size + sentence_len <= chunk_size:
            current_chunk.append(sentence)
            current_size += sentence_len
        else:
            # Save current chunk
            chunks.append(' '.join(current_chunk))
            
            # Start new chunk with overlap
            overlap_text = ' '.join(current_chunk)
            if len(overlap_text) > overlap:
                overlap_text = overlap_text[-overlap:]
            
            current_chunk = [overlap_text, sentence]
            current_size = len(overlap_text) + sentence_len
    
    # Add remaining chunk
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return [c.strip() for c in chunks if c.strip()]

def extract_docx_paragraphs_with_headings(file_bytes: bytes):
    """
    Returns a list of tuples: (paragraph_text, heading)
    """
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    current_heading = "No Heading"

    for p in doc.paragraphs:
        style = p.style.name
        if style.startswith("Heading"):
            current_heading = p.text.strip() or current_heading
        elif p.text.strip():
            paragraphs.append((p.text.strip(), current_heading))

    return paragraphs

def extract_pdf_paragraphs_with_headings(file_bytes: bytes):
    """
    Returns a list of tuples: (paragraph_text, heading)
    Improved heuristic: lines in ALL CAPS, short, or bold are headings
    """
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    paragraphs = []
    current_heading = "No Heading"

    for page in reader.pages:
        text = page.extract_text() or ""
        lines = text.split('\n')
        
        current_para = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_para:
                    para_text = ' '.join(current_para)
                    paragraphs.append((para_text, current_heading))
                    current_para = []
                continue
            
            # Improved heading detection
            is_heading = (
                (line.isupper() and len(line.split()) < 10) or  # ALL CAPS and short
                (line.endswith(':') and len(line.split()) < 15) or  # Ends with colon
                (len(line) < 60 and line[0].isupper() and not line.endswith('.'))  # Short, capitalized, no period
            )
            
            if is_heading:
                if current_para:
                    para_text = ' '.join(current_para)
                    paragraphs.append((para_text, current_heading))
                    current_para = []
                current_heading = line
            else:
                current_para.append(line)
        
        # Add remaining paragraph
        if current_para:
            para_text = ' '.join(current_para)
            paragraphs.append((para_text, current_heading))

    return paragraphs

def merge_paragraphs_into_chunks(paragraphs, max_chunk_size=CHUNK_SIZE_CHARS):
    """
    Merge paragraphs into semantically coherent chunks with metadata.
    Uses smart chunking to respect sentence boundaries.
    """
    chunks = []
    current_chunk = ""
    current_heading = ""
    start_index = 0

    for i, (para_text, para_heading) in enumerate(paragraphs):
        # Initialize first chunk
        if not current_chunk:
            current_chunk = para_text
            current_heading = para_heading
            start_index = i
        # Try to merge if within size limit
        elif len(current_chunk) + len(para_text) + 2 <= max_chunk_size:
            current_chunk += "\n\n" + para_text
        # Chunk is full, save it and start new one
        else:
            # Use smart chunking if chunk is very large
            if len(current_chunk) > max_chunk_size * 1.2:
                sub_chunks = smart_chunk_text(current_chunk, max_chunk_size, CHUNK_OVERLAP_CHARS)
                for sc in sub_chunks:
                    chunks.append({
                        "text": sc.strip(),
                        "heading": current_heading,
                        "paragraph_index": start_index
                    })
            else:
                chunks.append({
                    "text": current_chunk.strip(),
                    "heading": current_heading,
                    "paragraph_index": start_index
                })
            
            current_chunk = para_text
            current_heading = para_heading
            start_index = i

    # Add final chunk
    if current_chunk:
        if len(current_chunk) > max_chunk_size * 1.2:
            sub_chunks = smart_chunk_text(current_chunk, max_chunk_size, CHUNK_OVERLAP_CHARS)
            for sc in sub_chunks:
                chunks.append({
                    "text": sc.strip(),
                    "heading": current_heading,
                    "paragraph_index": start_index
                })
        else:
            chunks.append({
                "text": current_chunk.strip(),
                "heading": current_heading,
                "paragraph_index": start_index
            })

    return chunks

# ==========================================
# 4. DEFAULT SYSTEM PROMPT (Fallback Only)
# ==========================================
DEFAULT_SYSTEM_PROMPT = """You are an AI Co-Pilot for accessibility and inclusive design, specifically supporting Mekong Inclusive Ventures (MIV) practitioners, educators, and Entrepreneur Support Organizations (ESOs).

Provide clear, concise, and actionable advice based on the provided context.
Focus on accuracy, brevity, and professionalism.

Structure responses as:
- Direct answer first
- Step-by-step guidance when needed
- Relevant tool links or examples
- Bullet points for clarity
- Provide URL links for all relevant sources.

Do not use overly friendly or casual language like "I'd be happy to help", "Sure thing!", or excessive exclamation marks.

If the context does not contain the answer, say, "I don't have specific information on this in the MIV knowledge base, but here is general best practice," followed by helpful guidance."""

# ==========================================
# 5. FASTAPI APP & ROUTES
# ==========================================
app = FastAPI(title="MIV AI Co-Pilot API")

# CORS - Allowing all origins to prevent connection issues
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def home():
    return {"status": "online", "message": "MIV AI Co-Pilot Brain is running 🧠"}

# -----------------------
# Ingest Endpoint
# -----------------------
@app.post("/ingest", response_model=IngestResponse)
async def ingest_endpoint(file: UploadFile = File(...), target_index: str = "kb"):
    """
    Upload and ingest a file into the vector database.

    KB (default):
      - PDF
      - DOCX
      - TXT
      - Generic JSON {topic, content}

    KM (target_index="km"):
      - JSON Knowledge Map with text_to_embed + metadata
    """
    start_time = time.time()
    filename = file.filename

    # -----------------------------
    # SELECT INDEX
    # -----------------------------
    if target_index == "km":
        index_target = pc.Index(KNOWLEDGE_MAP_INDEX_NAME)
    else:
        index_target = pc.Index(PINECONE_INDEX_NAME)

    # -----------------------------
    # FILE TYPE VALIDATION
    # -----------------------------
    if not filename.lower().endswith(('.pdf', '.docx', '.json', '.txt')):
        raise HTTPException(
            status_code=400,
            detail="Invalid file type. Only PDF, DOCX, TXT, or JSON supported."
        )

    try:
        # Read file bytes first
        file_bytes = await file.read()
        paragraph_chunks = []

        # =========================================================
        # KB INGESTION (IMPROVED CHUNKING)
        # =========================================================
        if target_index != "km":
            print(f"📄 Processing KB file: {filename}")

            if filename.lower().endswith('.pdf'):
                print("  Extracting PDF paragraphs...")
                paragraphs = extract_pdf_paragraphs_with_headings(file_bytes)
                paragraph_chunks_raw = merge_paragraphs_into_chunks(paragraphs)

            elif filename.lower().endswith('.docx'):
                print("  Extracting DOCX paragraphs...")
                paragraphs = extract_docx_paragraphs_with_headings(file_bytes)
                paragraph_chunks_raw = merge_paragraphs_into_chunks(paragraphs)

            elif filename.lower().endswith('.txt'):
                print("  Extracting TXT content...")
                text = file_bytes.decode("utf-8", errors="ignore")
                # Use smart chunking for plain text
                chunks_text = smart_chunk_text(text, CHUNK_SIZE_CHARS, CHUNK_OVERLAP_CHARS)
                paragraph_chunks_raw = [
                    {"text": chunk, "heading": "No Heading", "paragraph_index": i} 
                    for i, chunk in enumerate(chunks_text)
                ]

            elif filename.lower().endswith('.json'):
                print("  Extracting JSON content...")
                data = json.loads(file_bytes)
                paragraph_chunks_raw = []
                for i, entry in enumerate(data):
                    content = entry.get("content", "")
                    # Apply smart chunking to JSON content if too large
                    if len(content) > CHUNK_SIZE_CHARS * 1.2:
                        chunks_text = smart_chunk_text(content, CHUNK_SIZE_CHARS, CHUNK_OVERLAP_CHARS)
                        for j, chunk in enumerate(chunks_text):
                            paragraph_chunks_raw.append({
                                "text": chunk,
                                "heading": entry.get("topic", "No Heading"),
                                "paragraph_index": f"{i}-{j}"
                            })
                    else:
                        paragraph_chunks_raw.append({
                            "text": content,
                            "heading": entry.get("topic", "No Heading"),
                            "paragraph_index": i
                        })

            # Convert paragraph_chunks_raw to paragraph_chunks
            for chunk in paragraph_chunks_raw:
                if chunk.get("text", "").strip():  # Only add non-empty chunks
                    paragraph_chunks.append(chunk)

            print(f"  ✅ Extracted {len(paragraph_chunks)} chunks (avg size: {sum(len(c['text']) for c in paragraph_chunks) // len(paragraph_chunks) if paragraph_chunks else 0} chars)")

        # =========================================================
        # KM INGESTION
        # =========================================================
        else:
            print(f"🗺️ Processing KM file: {filename}")
            km_data = json.loads(file_bytes)

            for i, entry in enumerate(km_data):
                text_for_embedding = entry.get("text_to_embed", "").strip()
                if not text_for_embedding:
                    print(f"  ⚠️ Skipping entry {i}: no text_to_embed")
                    continue

                paragraph_chunks.append({
                    "text": text_for_embedding,
                    "paragraph_index": i,
                    "metadata": {
                        "source": filename,
                        "user_intent": entry.get("user_intent", ""),
                        "tool_name": entry.get("tool_name", ""),
                        "url": entry.get("url") or entry.get("URL", ""),
                        "text_to_embed": text_for_embedding
                    }
                })

            print(f"  ✅ Extracted {len(paragraph_chunks)} KM entries")

        # Check if we have chunks before deleting
        if not paragraph_chunks:
            raise HTTPException(
                status_code=400,
                detail=f"No valid content extracted from {filename}. File may be empty or corrupted."
            )

        # -----------------------------
        # CLEAN EXISTING VECTORS (only after successful extraction)
        # -----------------------------
        print(f"🗑️ Deleting existing vectors for: {filename}")
        try:
            index_target.delete(filter={"source": filename})
            time.sleep(1)  # Wait for deletion to propagate
        except Exception as e:
            print(f"  ⚠️ Could not delete existing vectors: {e}")

        # =========================================================
        # EMBEDDING + UPSERT WITH DEDUPLICATION
        # =========================================================
        print(f"🔄 Embedding and upserting {len(paragraph_chunks)} chunks...")
        vectors_to_upsert = []
        seen_texts = set()  # Prevent duplicate chunks

        for chunk in paragraph_chunks:
            text = chunk["text"]
            para_idx = chunk["paragraph_index"]
            
            # Skip duplicates
            text_hash = hash(text)
            if text_hash in seen_texts:
                print(f"  ⏭️ Skipping duplicate chunk {para_idx}")
                continue
            seen_texts.add(text_hash)

            # Generate embedding
            try:
                embedding_response = client.models.embed_content(
                    model=EMBED_MODEL_NAME,
                    contents=text
                )
                vector = embedding_response.embeddings[0].values
            except Exception as e:
                print(f"  ❌ Error embedding chunk {para_idx}: {e}")
                continue

            # -----------------------------
            # KB UPSERT (CLEAN)
            # -----------------------------
            if target_index != "km":
                chunk_id = f"{filename}-para-{para_idx}"

                vectors_to_upsert.append((
                    chunk_id,
                    vector,
                    {
                        "text": text,
                        "heading": chunk.get("heading", "No Heading"),
                        "source": filename,
                        "chunk_size": len(text)  # Track chunk size for debugging
                    }
                ))

            # -----------------------------
            # KM UPSERT (INTENT-BASED)
            # -----------------------------
            else:
                metadata = chunk["metadata"]
                chunk_id = f"km-{filename}-{para_idx}"

                vectors_to_upsert.append((
                    chunk_id,
                    vector,
                    {
                        "text": metadata["text_to_embed"],
                        "source": metadata["source"],
                        "user_intent": metadata["user_intent"],
                        "tool_name": metadata["tool_name"],
                        "url": metadata["url"],
                    }
                ))

            # Batch upsert
            if len(vectors_to_upsert) >= BATCH_SIZE:
                index_target.upsert(vectors=vectors_to_upsert)
                print(f"  📤 Upserted batch of {len(vectors_to_upsert)} vectors")
                vectors_to_upsert = []

        # Final flush
        if vectors_to_upsert:
            index_target.upsert(vectors=vectors_to_upsert)
            print(f"  📤 Upserted final batch of {len(vectors_to_upsert)} vectors")

        elapsed = time.time() - start_time
        print(f"✅ Ingestion complete in {elapsed:.2f}s")

        return IngestResponse(
            success=True,
            message=f"Successfully ingested {filename}",
            chunks_added=len(paragraph_chunks),
            filename=filename
        )

    except json.JSONDecodeError as e:
        print(f"❌ JSON parsing error: {e}")
        raise HTTPException(
            status_code=400,
            detail=f"Invalid JSON file: {str(e)}"
        )
    except UnicodeDecodeError as e:
        print(f"❌ Text encoding error: {e}")
        raise HTTPException(
            status_code=400,
            detail=f"File encoding error. Please ensure file is UTF-8 encoded: {str(e)}"
        )
    except Exception as e:
        print(f"❌ Unexpected error during ingestion: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error processing file: {str(e)}"
        )

# -----------------------
# Chat Endpoint (Dual Index) - OPTIMIZED RETRIEVAL
# -----------------------
@app.post("/chat", response_model=ChatResponse)
async def chat_endpoint(req: ChatRequest):
    start_time = time.time()
    
    # 1️⃣ Get user question
    question = req.query.strip()
    print(f"\n🔥 Received Question: {question}")

    # Use passed system prompt or fall back to default
    system_prompt = req.system_prompt if req.system_prompt else DEFAULT_SYSTEM_PROMPT
    print(f"📋 Using System Prompt: {system_prompt[:100]}...")

    # 2️⃣ Handle greetings first
    greetings = ['hi', 'hello', 'hey', 'good morning', 'good afternoon']
    if question.lower() in greetings:
        return {
            "response": "Hello! 👋 I'm your AI Co-Pilot for accessibility. I can help you find tools, understand guidelines, or improve your content. What would you like to know?",
            "sources": []
        }

    # 3️⃣ Detect formatting instructions
    formatted_question = question  # default
    q_lower = question.lower()

    if "in three words" in q_lower:
        formatted_question += " (Answer in exactly three words, separated by commas, no extra commentary)"
    elif "bullet points" in q_lower:
        formatted_question += (
            " (Answer as bullet points: each tip on a separate line starting with '-', do not combine multiple tips in one paragraph, no inline asterisks)"
        )
    elif "numbered steps" in q_lower:
        formatted_question += (
            " (Answer as numbered steps: each step on a separate line starting with its number, no extra commentary)"
        )

    try:
        # --- EMBED USER QUESTION ---
        embedding_response = client.models.embed_content(
            model=EMBED_MODEL_NAME,
            contents=question
        )
        query_embedding = embedding_response.embeddings[0].values

        # --- STEP 1: QUERY KNOWLEDGE MAP ---
        km_results = index_km.query(
            vector=query_embedding,
            top_k=2,  # Increased from 1 to get better coverage
            include_metadata=True
        )

        km_text = ""
        km_topic = question
        
        if km_results['matches']:
            # Use the best match for topic expansion
            km_text = km_results['matches'][0]['metadata'].get('text', '')
            km_topic = km_text
            print("🔹 KM Retrieved:")
            for match in km_results['matches']:
                metadata = match.get('metadata', {})
                print(f"  - User Intent: {metadata.get('user_intent')}")
                print(f"  - Source: {metadata.get('source')}")
                print(f"  - Score: {match['score']:.3f}")
                print(f"  - Text preview: {metadata.get('text', '')[:150]}")

        # --- STEP 2: QUERY KNOWLEDGE BASE using KM topic ---
        kb_embedding_response = client.models.embed_content(
            model=EMBED_MODEL_NAME,
            contents=km_topic
        )
        kb_query_embedding = kb_embedding_response.embeddings[0].values

        kb_results = index_kb.query(
            vector=kb_query_embedding,
            top_k=req.top_k,
            include_metadata=True
        )

        print("🔹 KB Retrieved:")
        for match in kb_results['matches']:
            metadata = match.get('metadata', {})
            print(f"  - Source: {metadata.get('source')}")
            print(f"  - Heading: {metadata.get('heading')}")
            print(f"  - Score: {match['score']:.3f}")
            print(f"  - Chunk size: {metadata.get('chunk_size', 'unknown')} chars")
            print(f"  - Text preview: {metadata.get('text', '')[:150]}")  
            
        # --- BUILD CONTEXT WITH DEDUPLICATION ---
        retrieved_chunks = []
        context_text_list = []
        seen_sources = set()

        # Add Knowledge Map snippet first
        if km_text:
            context_text_list.append(f"[Source: Knowledge Map]\n{km_text}")
            retrieved_chunks.append({
                "text": km_text[:200]+"...", 
                "source": "Knowledge Map", 
                "score": 1.0
            })

        # Add Knowledge Base results (filter by relevance score)
        RELEVANCE_THRESHOLD = 0.7  # Only include chunks with score > 0.7
        
        for match in kb_results['matches']:
            if match['score'] < RELEVANCE_THRESHOLD:
                print(f"  ⏭️ Skipping low relevance chunk (score: {match['score']:.3f})")
                continue
                
            metadata = match.get('metadata', {})
            text_content = metadata.get('text', '')
            source_name = metadata.get('source', 'Knowledge Base')
            
            # Prevent duplicate sources from dominating context
            source_key = f"{source_name}:{text_content[:50]}"
            if source_key in seen_sources:
                continue
            seen_sources.add(source_key)

            context_text_list.append(f"[Source: {source_name}]\n{text_content}")
            retrieved_chunks.append({
                "text": text_content[:200]+"...", 
                "source": source_name, 
                "score": match['score']
            })

        full_context = "\n\n---\n\n".join(context_text_list)
        
        # Limit total context size to prevent token overflow
        MAX_CONTEXT_CHARS = 8000  # ~2000 tokens for context
        if len(full_context) > MAX_CONTEXT_CHARS:
            full_context = full_context[:MAX_CONTEXT_CHARS] + "\n\n[Context truncated...]"
            print(f"⚠️ Context truncated to {MAX_CONTEXT_CHARS} chars")

        # --- GENERATE AI RESPONSE USING PASSED SYSTEM PROMPT ---
        prompt = f"""{system_prompt}

CONTEXT FROM KNOWLEDGE MAP + KNOWLEDGE BASE:
{full_context}

USER QUESTION:
{formatted_question}
"""
        response = client.models.generate_content(
            model=CHAT_MODEL_NAME,
            contents=prompt
        )

        elapsed = time.time() - start_time
        print(f"✅ Reply generated in {elapsed:.2f}s")

        return {"response": response.text, "sources": retrieved_chunks}

    except Exception as e:
        print(f"❌ Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# -----------------------
# List Documents Endpoint
# -----------------------
@app.get("/list-documents")
async def list_documents():
    try:
        results = index_kb.query(
            vector=[0.0] * 768,
            top_k=1000,
            include_metadata=True
        )
        sources = set()
        for match in results.get('matches', []):
            metadata = match.get('metadata', {})
            if 'source' in metadata:
                sources.add(metadata['source'])
        documents = [{"filename": src} for src in sorted(sources)]
        return {"success": True, "documents": documents, "total_chunks_sampled": len(results.get('matches', []))}
    except Exception as e:
        print(f"❌ Error listing documents: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# -----------------------
# List Knowledge Maps Endpoint
# ----------------------- 
@app.get("/list-knowledge-maps")
async def list_km():
    try:
        results = index_km.query(vector=[0.0]*768, top_k=1000, include_metadata=True)
        sources = set()
        for match in results.get('matches', []):
            metadata = match.get('metadata', {})
            if 'source' in metadata:
                sources.add(metadata['source'])
        km_docs = [{"filename": src} for src in sorted(sources)]
        return {"success": True, "knowledge_maps": km_docs}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))